#!/usr/bin/env python3
"""
Enhanced document service with database integration, caching, and advanced features
"""

import os
import uuid
import asyncio
from typing import List, Optional, Dict, Any
from fastapi import UploadFile, HTTPException
import logging
from datetime import datetime
import mimetypes
import hashlib

# Document processing imports
import fitz  # PyMuPDF
from docx import Document as DocxDocument
from langdetect import detect

# Internal services
from models import Document, Language
from services.database_service import DatabaseService
from services.cache_service import cache_service, generate_cache_key
from services.secret_manager import secret_manager

logger = logging.getLogger(__name__)

class EnhancedDocumentService:
    """Enhanced document service with database, caching, and advanced features"""
    
    def __init__(self):
        self.database = DatabaseService()
        self.max_file_size = 50 * 1024 * 1024  # 50MB
        self.supported_types = {
            'pdf': 'application/pdf',
            'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'doc': 'application/msword',
            'txt': 'text/plain',
            'md': 'text/markdown',
            'rtf': 'application/rtf'
        }
        
    async def initialize(self):
        """Initialize the service"""
        await self.database.initialize()
        logger.info("Enhanced document service initialized")
    
    async def upload_and_process_document(
        self,
        file: UploadFile,
        user_id: Optional[str] = None
    ) -> Document:
        """Upload and process a document with validation and optimization"""
        
        # Validate file
        await self._validate_file(file)
        
        # Generate document ID
        document_id = str(uuid.uuid4())
        
        # Create temporary file
        temp_file_path = f"/tmp/{document_id}_{file.filename}"
        
        try:
            # Save uploaded file
            with open(temp_file_path, "wb") as temp_file:
                content = await file.read()
                temp_file.write(content)
            
            # Process document
            document = await self._process_document_file(
                temp_file_path,
                document_id,
                file.filename,
                user_id
            )
            
            # Save to database
            success = await self.database.save_document(document)
            if not success:
                raise HTTPException(status_code=500, detail="Failed to save document")
            
            # Cache document for quick access
            await self._cache_document(document)
            
            logger.info(f"Document processed successfully: {document_id}")
            return document
            
        except Exception as e:
            logger.error(f"Error processing document: {str(e)}")
            raise HTTPException(status_code=500, detail=f"Document processing failed: {str(e)}")
        
        finally:
            # Clean up temporary file
            if os.path.exists(temp_file_path):
                os.remove(temp_file_path)
    
    async def get_document(self, document_id: str, user_id: Optional[str] = None) -> Optional[Document]:
        """Get document with caching and user validation"""
        
        # Try cache first
        cache_key = generate_cache_key("document", document_id)
        cached_doc = await cache_service.get(cache_key)
        
        if cached_doc:
            document = Document(**cached_doc)
            # Validate user access
            if user_id and document.user_id != user_id:
                return None
            return document
        
        # Get from database
        document = await self.database.get_document(document_id)
        
        if document:
            # Validate user access
            if user_id and document.user_id != user_id:
                return None
            
            # Cache for future requests
            await self._cache_document(document)
            
        return document
    
    async def list_documents(
        self,
        user_id: Optional[str] = None,
        limit: int = 100,
        offset: int = 0
    ) -> List[Document]:
        """List documents with pagination and caching"""
        
        # Try cache for user-specific lists
        if user_id:
            cache_key = generate_cache_key("user_documents", user_id, limit, offset)
            cached_docs = await cache_service.get(cache_key)
            if cached_docs:
                return [Document(**doc) for doc in cached_docs]
        
        # Get from database
        documents = await self.database.list_documents(user_id, limit)
        
        # Apply offset (database service should handle this, but fallback)
        if offset > 0:
            documents = documents[offset:]
        
        # Cache user-specific lists
        if user_id and documents:
            cache_key = generate_cache_key("user_documents", user_id, limit, offset)
            doc_dicts = [doc.dict() for doc in documents]
            await cache_service.set(cache_key, doc_dicts, ttl=300)  # 5 minutes
        
        return documents
    
    async def search_documents(
        self,
        query: str,
        user_id: Optional[str] = None,
        limit: int = 50
    ) -> List[Document]:
        """Search documents with caching"""
        
        # Cache search results
        cache_key = generate_cache_key("search", query, user_id, limit)
        cached_results = await cache_service.get(cache_key)
        
        if cached_results:
            return [Document(**doc) for doc in cached_results]
        
        # Search in database
        documents = await self.database.search_documents(query, user_id)
        
        # Limit results
        if len(documents) > limit:
            documents = documents[:limit]
        
        # Cache results
        if documents:
            doc_dicts = [doc.dict() for doc in documents]
            await cache_service.set(cache_key, doc_dicts, ttl=600)  # 10 minutes
        
        return documents
    
    async def delete_document(self, document_id: str, user_id: Optional[str] = None) -> bool:
        """Delete document with user validation and cache invalidation"""
        
        # Get document to validate ownership
        document = await self.get_document(document_id, user_id)
        if not document:
            return False
        
        # Validate user can delete
        if user_id and document.user_id != user_id:
            raise HTTPException(status_code=403, detail="Not authorized to delete this document")
        
        # Delete from database
        success = await self.database.delete_document(document_id)
        
        if success:
            # Invalidate caches
            await self._invalidate_document_caches(document_id, document.user_id)
            logger.info(f"Document deleted: {document_id}")
        
        return success
    
    async def update_document(self, document: Document) -> bool:
        """Update document with cache invalidation"""
        
        success = await self.database.update_document(document)
        
        if success:
            # Update cache
            await self._cache_document(document)
            # Invalidate related caches
            await self._invalidate_document_caches(document.document_id, document.user_id)
        
        return success
    
    async def get_document_statistics(self, user_id: Optional[str] = None) -> Dict[str, Any]:
        """Get document statistics"""
        
        cache_key = generate_cache_key("doc_stats", user_id or "global")
        cached_stats = await cache_service.get(cache_key)
        
        if cached_stats:
            return cached_stats
        
        # Get documents
        documents = await self.list_documents(user_id, limit=10000)  # Get all for stats
        
        # Calculate statistics
        stats = {
            "total_documents": len(documents),
            "total_size": sum(doc.meta.get("file_size", 0) for doc in documents),
            "file_types": {},
            "languages": {},
            "recent_uploads": 0
        }
        
        # Count by file type and language
        for doc in documents:
            file_type = doc.meta.get("file_type", "unknown")
            stats["file_types"][file_type] = stats["file_types"].get(file_type, 0) + 1
            
            language = doc.detected_language or "unknown"
            stats["languages"][language] = stats["languages"].get(language, 0) + 1
            
            # Count recent uploads (last 24 hours)
            if doc.created_at:
                try:
                    created_time = datetime.fromisoformat(doc.created_at.replace('Z', '+00:00'))
                    if (datetime.now(created_time.tzinfo) - created_time).days < 1:
                        stats["recent_uploads"] += 1
                except:
                    pass
        
        # Cache statistics
        await cache_service.set(cache_key, stats, ttl=1800)  # 30 minutes
        
        return stats
    
    async def _validate_file(self, file: UploadFile):
        """Validate uploaded file"""
        
        # Check file size
        if hasattr(file, 'size') and file.size > self.max_file_size:
            raise HTTPException(
                status_code=413,
                detail=f"File too large. Maximum size is {self.max_file_size // (1024*1024)}MB"
            )
        
        # Check file type
        file_extension = file.filename.lower().split('.')[-1] if '.' in file.filename else ''
        if file_extension not in self.supported_types:
            raise HTTPException(
                status_code=415,
                detail=f"Unsupported file type. Supported types: {list(self.supported_types.keys())}"
            )
        
        # Validate filename
        if not file.filename or len(file.filename) > 255:
            raise HTTPException(status_code=400, detail="Invalid filename")
    
    async def _process_document_file(
        self,
        file_path: str,
        document_id: str,
        filename: str,
        user_id: Optional[str] = None
    ) -> Document:
        """Process document file and extract content"""
        
        # Extract text content
        content = await self._extract_text_from_file(file_path, filename)
        
        # Detect language
        detected_language = self._detect_language(content)
        
        # Extract title
        title = self._extract_title(content, filename)
        
        # Get file metadata
        file_stats = os.stat(file_path)
        
        # Create document
        document = Document(
            document_id=document_id,
            filename=filename,
            title=title,
            text=content,
            user_id=user_id,
            detected_language=detected_language,
            meta={
                "file_size": file_stats.st_size,
                "file_type": self._get_file_type(filename),
                "content_hash": self._calculate_content_hash(content),
                "processing_time": datetime.utcnow().isoformat()
            }
        )
        
        return document
    
    async def _extract_text_from_file(self, file_path: str, filename: str) -> str:
        """Extract text from file based on type"""
        
        file_extension = filename.lower().split('.')[-1]
        
        try:
            if file_extension == 'pdf':
                return await self._extract_text_from_pdf(file_path)
            elif file_extension in ['docx', 'doc']:
                return await self._extract_text_from_docx(file_path)
            elif file_extension in ['txt', 'md']:
                return await self._extract_text_from_txt(file_path)
            else:
                # Try as text file
                return await self._extract_text_from_txt(file_path)
                
        except Exception as e:
            logger.error(f"Error extracting text from {filename}: {str(e)}")
            raise HTTPException(status_code=422, detail=f"Could not extract text from file: {str(e)}")
    
    async def _extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF using PyMuPDF"""
        def extract():
            doc = fitz.open(file_path)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text.strip()
        
        # Run in thread pool to avoid blocking
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, extract)
    
    async def _extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        def extract():
            doc = DocxDocument(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text.strip()
        
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, extract)
    
    async def _extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from text file"""
        def extract():
            try:
                with open(file_path, "r", encoding="utf-8") as f:
                    return f.read().strip()
            except UnicodeDecodeError:
                with open(file_path, "r", encoding="latin-1") as f:
                    return f.read().strip()
        
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, extract)
    
    def _detect_language(self, text: str) -> str:
        """Detect language of text"""
        try:
            lang_code = detect(text)
            # Map to supported languages
            lang_map = {
                "en": "en", "te": "te", "hi": "hi",
                "es": "es", "fr": "fr", "de": "de"
            }
            return lang_map.get(lang_code, "en")
        except:
            return "en"
    
    def _extract_title(self, content: str, filename: str) -> str:
        """Extract title from content or filename"""
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if line and len(line) < 100 and len(line) > 3:
                return line
        
        # Fallback to filename without extension
        return filename.rsplit('.', 1)[0]
    
    def _get_file_type(self, filename: str) -> str:
        """Get file type from filename"""
        return filename.lower().split('.')[-1] if '.' in filename else 'unknown'
    
    def _calculate_content_hash(self, content: str) -> str:
        """Calculate hash of content for deduplication"""
        return hashlib.sha256(content.encode()).hexdigest()
    
    async def _cache_document(self, document: Document):
        """Cache document for quick access"""
        cache_key = generate_cache_key("document", document.document_id)
        await cache_service.set(cache_key, document.dict(), ttl=3600)  # 1 hour
    
    async def _invalidate_document_caches(self, document_id: str, user_id: Optional[str]):
        """Invalidate related caches when document changes"""
        
        # Invalidate document cache
        cache_key = generate_cache_key("document", document_id)
        await cache_service.delete(cache_key)
        
        # Invalidate user document lists (approximate)
        if user_id:
            # This is a simple approach - in production, use cache tags
            for limit in [10, 50, 100]:
                for offset in [0, 10, 50]:
                    cache_key = generate_cache_key("user_documents", user_id, limit, offset)
                    await cache_service.delete(cache_key)
        
        # Invalidate statistics
        stats_key = generate_cache_key("doc_stats", user_id or "global")
        await cache_service.delete(stats_key)

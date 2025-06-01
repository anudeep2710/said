import os
import fitz  # PyMuPDF
from docx import Document as DocxDocument
from typing import List, Optional, Union
from models import Document, Language
from services.storage_service import StorageService
from services.mock_storage_service import MockStorageService

class DocumentService:
    """Service for handling document operations"""

    def __init__(self, storage_service: Optional[Union[StorageService, MockStorageService]] = None):
        # Use provided storage service or default to mock storage
        if storage_service is None:
            self.storage_service = MockStorageService()
        else:
            self.storage_service = storage_service

    async def process_document(self, file_path: str, document_id: str, filename: str, user_id: Optional[str] = None) -> Document:
        """Process and store a document"""
        try:
            # Extract text based on file type
            content = self._extract_text_from_file(file_path, filename)

            # Detect language
            detected_language = self._detect_language(content)

            # Extract title from content (first line or filename)
            title = self._extract_title(content, filename)

            # Create document object
            document = Document(
                document_id=document_id,
                filename=filename,
                title=title,
                text=content,
                user_id=user_id,
                detected_language=detected_language,
                meta={
                    "file_size": os.path.getsize(file_path),
                    "file_type": self._get_file_type(filename)
                }
            )

            # Store document
            await self.storage_service.store_document(document, file_path)

            return document

        except Exception as e:
            raise Exception(f"Error processing document: {str(e)}")

    def _extract_text_from_file(self, file_path: str, filename: str) -> str:
        """Extract text from different file types"""
        file_extension = filename.lower().split('.')[-1]

        if file_extension == 'pdf':
            return self._extract_text_from_pdf(file_path)
        elif file_extension in ['docx', 'doc']:
            return self._extract_text_from_docx(file_path)
        elif file_extension == 'txt':
            return self._extract_text_from_txt(file_path)
        else:
            # Try to read as text file
            try:
                return self._extract_text_from_txt(file_path)
            except:
                raise Exception(f"Unsupported file type: {file_extension}")

    def _extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF using PyMuPDF"""
        try:
            doc = fitz.open(file_path)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text.strip()
        except Exception as e:
            raise Exception(f"Error extracting text from PDF: {str(e)}")

    def _extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = DocxDocument(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text.strip()
        except Exception as e:
            raise Exception(f"Error extracting text from DOCX: {str(e)}")

    def _extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read().strip()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, "r", encoding="latin-1") as f:
                return f.read().strip()

    def _detect_language(self, text: str) -> Language:
        """Detect language of text"""
        try:
            from langdetect import detect
            lang_code = detect(text)
            # Map language codes to our Language enum
            lang_map = {
                "en": Language.ENGLISH,
                "te": Language.TELUGU,
                "hi": Language.HINDI,
                "es": Language.SPANISH,
                "fr": Language.FRENCH,
                "de": Language.GERMAN
            }
            return lang_map.get(lang_code, Language.ENGLISH)
        except:
            return Language.ENGLISH

    def _extract_title(self, content: str, filename: str) -> str:
        """Extract title from content or use filename"""
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if line and len(line) < 100:  # Reasonable title length
                return line
        # Fallback to filename without extension
        return filename.rsplit('.', 1)[0]

    def _get_file_type(self, filename: str) -> str:
        """Get file type from filename"""
        return filename.lower().split('.')[-1] if '.' in filename else 'unknown'

    async def get_document(self, document_id: str) -> Optional[Document]:
        """Get document by ID"""
        try:
            return await self.storage_service.retrieve_document(document_id)
        except Exception as e:
            raise Exception(f"Error getting document: {str(e)}")

    async def list_documents(self, user_id: Optional[str] = None) -> List[Document]:
        """List documents (optionally filtered by user)"""
        try:
            all_documents = await self.storage_service.list_documents()
            if user_id:
                # Filter documents by user_id
                return [doc for doc in all_documents if doc.user_id == user_id]
            return all_documents
        except Exception as e:
            raise Exception(f"Error listing documents: {str(e)}")

    async def delete_document(self, document_id: str) -> None:
        """Delete a document"""
        try:
            await self.storage_service.delete_document(document_id)
        except Exception as e:
            raise Exception(f"Error deleting document: {str(e)}")

    async def get_document_content(self, document_id: str) -> str:
        """Get document content"""
        try:
            return await self.storage_service.get_document_content(document_id)
        except Exception as e:
            raise Exception(f"Error getting document content: {str(e)}")